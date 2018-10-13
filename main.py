import os
import string as s
import bs4 as bs
import urllib.request
import docx2txt
import tkinter
import googleapiclient.errors
import googleapiclient.discovery
import regex as re
from tkinter.filedialog import askopenfilename
from docx import Document
from docx.shared import Pt


def google_search(query, api_key, cse_id, **kwargs):
    """
    Uses Google API to query a search and returns the search results. Requires valid API key and CSE ID.

    :param query: Query, passed as string.
    :param api_key: Google API Key, passed as string.
    :param cse_id: Google CSE ID, passed as string.
    :param kwargs: Arguments for Google API list() function.
    """

    try:
        service = googleapiclient.discovery.build("customsearch", "v1", developerKey=api_key)
        res = service.cse().list(q=query, cx=cse_id, **kwargs).execute()
        return res['items'][0]
    except googleapiclient.errors.HttpError:
        exit("Missing Google API Key and/or CSE ID.")


def scrape(terms, api_key, cse_id):
    """
    Google a definition for each term provided and return them.

    :param terms: List of terms to search using Google API.
    :param api_key: Google API Key, passed as string.
    :param cse_id: Google CSE ID, passed as string.
    :rtype: dict
    """
    # Clean up terms for Google Search. Removes all parentheses and text between parentheses.
    clean_terms = [re.sub(r'\([^)]*\)', '', term) for term in terms]
    not_found = []
    definitions = {}
    for i, clean_term in enumerate(clean_terms):
        print("Googling " + clean_term + "...")
        term = terms[i]
        try:
            # Try Investopedia.
            results = google_search(clean_term + " definition economics investopedia",
                                    num=1, api_key=api_key, cse_id=cse_id)
            if 'investopedia' in results['formattedUrl']:
                definitions[term] = results['pagemap']['metatags'][0]['twitter:description']
                continue
            # Now try Wikipedia.
            results = google_search(clean_term + " definition economics wikipedia",
                                    num=1, api_key=api_key, cse_id=cse_id)
            if 'wikipedia' in results['formattedUrl']:
                # Google shortens their website descriptions with ellipses if they are too long.
                # Wikipedia's descriptions are almost always shortened, so we visit Wikipedia
                # and grab the rest of the text snippet directly from the site html.
                definitions[term] = parse(results['formattedUrl'], results['snippet'])
                print(definitions[term])
                continue
        except Exception as e:
            print(e)
            not_found.append(term)

    return definitions, not_found


def parse(term_url, snippet):
    """
    Parses HTML text to grab the full paragraph containing the shortened snippet.
    Helper function for scrape().

    :param term_url: Page URL from Google API.
    :param snippet: Shortened snippet from Google API.
    :return: Full-length snippet grabbed from URL.
    """

    # Remove ellipses from snippet.
    clean_snippet = re.sub("[...]", "", snippet)
    # Create a fake user agent to bypass PollEvBot detection.
    url_headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                   '(KHTML, like Gecko) Chrome/61.0.3163.100 Safari/537.36'}
    # Grab and parse website HTML with BeautifulSoup.
    r = urllib.request.urlopen(urllib.request.Request(term_url, headers=url_headers))
    soup = bs.BeautifulSoup(r, "lxml")
    for paragraph in soup.body.find_all('p'):
        # If the snippet is part of the paragraph, return the whole paragraph.
        if clean_snippet in paragraph.text:
            return paragraph.text


def read_doc(filepath):
    """
    Finds and reads a docx file containing search terms.

    :param filepath: Path for docx file containing search terms.
    :return: List of terms within document.
    """
    allowed_chars = s.ascii_letters + s.digits + '-() '
    not_allowed_words = ["Vocab", "Chapter"]
    text = docx2txt.process(filepath)
    # Read the text, with newlines as delimiters.
    text = list(filter(None, text.split('\n')))
    terms = []
    for string in text:
        # If every character is allowed and the string doesn't contain the words "Vocab" or "Chapter",
        # it's probably a vocabulary term
        if all(char in allowed_chars for char in string) and not any(word in string for word in not_allowed_words):
            terms.append(string)
    return terms


def add_par(docname, text, font='Times New Roman', font_size=12, bold_text=False, underline_text=False):
    """
    Helper function for writing to docx documents.

    :param docname: Name of docx document to write to.
    :param text: Text to write, passed as a string.
    :param font: Type of font to use (Times New Roman, Calibri, etc.)
    :param font_size: Size of font to use.
    :param bold_text: Use bold text if true.
    :param underline_text: Underline text if true.
    """

    run = docname.add_paragraph().add_run(text)
    paragraph = run.font

    paragraph.name = font
    paragraph.size = Pt(font_size)
    paragraph.bold = bold_text
    paragraph.underline = underline_text


def add_mla_header(doc, *args):
    """
    Adds an MLA header to a docx document.

    :param doc: Docx document object.
    :param args: Parts for MLA header.
    """
    for arg in args:
        add_par(doc, arg)
    # Add an empty line after writing MLA header.
    add_par(doc, "")


def make_doc(docname, filepath, definitions, not_found, *args):
    """
    Creates a Word document with an MLA header, terms, definitions, and words we didn't find.

    :param docname: Name of document.
    :param filepath: Path (location) to save document.
    :param definitions: Dict with terms as keys and definitions as values.
    :param not_found: List of terms we didn't find.
    :param args: Parts for MLA header.
    """

    os.chdir(filepath)
    doc = Document()
    # My preferred MLA formatting. Can be modified as needed.
    print(*args)
    add_mla_header(doc, *args)
    add_par(doc, "Vocabulary List:", bold_text=True)
    # Write terms and definitions.
    for i, term in enumerate(definitions.keys()):
        add_par(doc, str(i+1) + ".   " + term)
        add_par(doc, definitions[term])
    # Write terms we didn't find definitions for.
    if not_found:
        add_par(doc, "Words not found:", underline_text=True)
        for term in not_found:
            add_par(doc, term)
    doc.save(docname)


def main(makefile, *args):
    """
    Reads a Word document containing terms without definitions.
    Googles the terms, then creates a new Word document with terms and definitions.
    """
    # Key and ID for Google's customsearch API.
    Key = ""
    ID = ""
    # Opens an instance of tkinter for user file selection.
    tkinter.Tk().withdraw()  # Close the root window
    vocab_fpath = askopenfilename()
    # Make sure the user selected a file.
    if vocab_fpath == '':
        exit("No input file selected.")
    print("File received.")
    terms = read_doc(vocab_fpath)
    definitions, not_found = scrape(terms, api_key=Key, cse_id=ID)
    make_doc(makefile, os.path.dirname(vocab_fpath), definitions, not_found, *args)
    os.startfile(makefile)


if __name__ == '__main__':
    main("Chapter 17 Vocabulary.docx", "Daniel Qiang", "2/4/18", "Sherman Per. 6", "Macroeconomic Objectives")



