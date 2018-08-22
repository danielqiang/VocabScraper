import bs4 as bs
import urllib.request
import docx2txt
import regex as re
from tkinter.filedialog import askopenfilename
from docx import Document
from docx.shared import Pt
from googleapiclient.discovery import build
import tkinter as tk
import os
import string as s
import time


# Key and ID for Google's customsearch API.
Key = "AIzaSyBLyswP6EvROXHsDo4O_BIQwrEx-Gp5GW4"
ID = "010141786122130001585:gob4xioziac"


def google_search(query, api_key=Key, cse_id=ID, **kwargs):
    """
    Uses Google API to query a search. Requires valid API key and CSE ID.

    :param query: Query, passed as string.
    :param api_key: Google API Key, passed as string.
    :param cse_id: Google CSE ID, passed as string.
    :param kwargs: Arguments for Google API list() function.
    :return: Dictionary containing search results.
    """

    service = build("customsearch", "v1", developerKey=api_key)
    res = service.cse().list(q=query, cx=cse_id, **kwargs).execute()
    return res['items'][0]


def scrape(terms):
    """
    For each term in a list of terms, query Google and return their definitions.
    This function is currently optimized for IB Economics terms, and can be modified as needed.

    :param terms: List of terms to search using Google API.
    :return: Dict containing terms as keys and definitions as values.
    """

    not_found = []
    definitions = {}

    # Clean up terms for Google Search. Removes all parentheses and text between parentheses.
    clean_terms = [re.sub(r'\([^)]*\)', '', term) for term in terms]

    for clean_term in clean_terms:
        print("Googling " + clean_term + "...")
        term = terms[clean_terms.index(clean_term)]
        try:
            # Try Investopedia.
            results = google_search(clean_term + " definition economics investopedia", num=1)
            # Step through API search results dict to obtain Google's featured snippet.
            # For Investopedia, this is stored under the tag 'twitter:description'.
            if 'investopedia' in results['formattedUrl']:
                definitions[term] = results['pagemap']['metatags'][0]['twitter:description']
                continue
            # Now try Wikipedia.
            results = google_search(clean_term + " definition economics wikipedia", num=1)
            if 'wikipedia' in results['formattedUrl']:
                # Wikipedia's snippet is automatically shortened by Google's API.
                # By cleaning the snippet and going to the page URL, we can find
                # the paragraph containing the shortened snippet and return the
                # entire paragraph as the definition.
                definitions[term] = parse(results['formattedUrl'], results['snippet'])
                print(definitions[term])
                continue
        # If there's an error, print the error and continue.
        except Exception as e:
            print(e)
            not_found.append(term)

    return definitions, not_found


def parse(term_url, snippet):
    """
    Parses HTML text to grab the full paragraph containing the shortened snippet.
    Helper function for scrape().

    :param term_url: Page URL from Google API.
    :param snippet: Shortened snippet from Google API.
    :return: Full-length snippet grabbed from URL.
    """

    # Remove ellipses from snippet.
    clean_snippet = re.sub("[...]", "", snippet)
    # Create a fake user agent to bypass bot detection.
    url_headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                   '(KHTML, like Gecko) Chrome/61.0.3163.100 Safari/537.36'}
    # Grab and parse website HTML with BeautifulSoup.
    req = urllib.request.Request(term_url, headers=url_headers)
    sauce = urllib.request.urlopen(req)
    soup = bs.BeautifulSoup(sauce, "lxml")
    for paragraph in soup.body.find_all('p'):
        # If the snippet is part of the paragraph, return the whole paragraph.
        if clean_snippet in paragraph.text:
            return paragraph.text


def read_doc(filepath):
    """
    Finds and reads a docx file containing search terms.

    :param filepath: Path for docx file containing search terms.
    :return: List of terms within document.
    """
    # Two lists - Chars allowed in terms and words not allowed in terms
    allowed_chars = s.ascii_letters + s.digits + '-() '
    not_allowed_words = ["Vocab", "Chapter"]

    text = docx2txt.process(filepath)
    # Split text with newlines as delimiters.
    text = list(filter(None, text.split('\n')))

    terms = []
    for string in text:
        # If every char in string is allowed and string contains no not-allowed words, append to terms
        if all(char in allowed_chars for char in string) and not any(word in string for word in not_allowed_words):
            terms.append(string)
    return terms


def add_par(docname, text, fname='Times New Roman', fsize=12, fbold=False, funderline=False):
    """
    Helper function for writing to docx documents.

    :param docname: Name of docx document to write to.
    :param text: Text to write, passed as a string.
    :param fname: Name of font to use.
    :param fsize: Size of font to use.
    :param fbold: Use bold text if true.
    :param funderline: Underline text if true.
    """

    run = docname.add_paragraph().add_run(text)
    font = run.font

    font.name = fname
    font.size = Pt(fsize)
    font.bold = fbold
    font.underline = funderline


def add_mla_header(doc, *args):
    """
    Adds an MLA header to a docx document.

    :param doc: Docx document object.
    :param args: Parts for MLA header.
    """
    for arg in args:
        add_par(doc, arg)
    # Add an empty line after writing MLA header.
    add_par(doc, "")


def make_doc(docname, filepath, definitions, not_found, *args):
    """
    Creates a Word document with an MLA header, terms, definitions, and words we didn't find.

    :param docname: Name of document.
    :param filepath: Path (location) to save document.
    :param definitions: Dict with terms as keys and definitions as values.
    :param not_found: List of terms we didn't find.
    :param args: Parts for MLA header.
    """

    os.chdir(filepath)
    doc = Document()
    # My preferred MLA formatting. Can be modified as needed.
    add_mla_header(doc, *args)
    add_par(doc, "Vocabulary List:", fbold=True)
    # Write terms and definitions.
    for i, term in enumerate(definitions.keys()):
        add_par(doc, str(i+1) + ".   " + term)
        add_par(doc, definitions[term])
    # Write terms we didn't find definitions for.
    if not_found:
        add_par(doc, "Words not found:", funderline=True)
        for term in not_found:
            add_par(doc, term)
    doc.save(docname)


def main(makefile, *args):
    """
    Reads a Word document containing terms without definitions.
    Googles the terms, then creates a new Word document with terms and definitions.
    """

    # Opens an instance of tkinter for user file selection.
    tk.Tk().withdraw()  # Close the root window
    vocab_fpath = askopenfilename()
    # Make sure the user selected a file.
    if vocab_fpath == '':
        exit("No input file selected.")
    print("File received.")

    start = time.time()

    # Grab a list of terms from Word Doc (ie. Econ Vocab List).
    terms = read_doc(vocab_fpath)
    # Scrape Google for term definitions.
    definitions, not_found = scrape(terms)
    # Make the Vocab Word Doc.
    make_doc(makefile, os.path.dirname(vocab_fpath), definitions, not_found, *args)

    end = time.time()
    print("Done! Googled " + str(len(terms)) + " terms in " + str(end - start) + " seconds.")
    # Open the Vocab Word Doc.
    os.startfile(makefile)


if __name__ == '__main__':
    main("Chapter 16 Vocabulary.docx", "Daniel Qiang", "2/4/18", "Sherman Per. 6", "Macroeconomic Objectives")



